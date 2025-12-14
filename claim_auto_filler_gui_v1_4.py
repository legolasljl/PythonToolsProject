import os
import re
import threading
import tkinter
from datetime import datetime, timedelta
from typing import Dict, Optional, Callable

import pandas as pd
import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- 业务逻辑类 (后端) ---
class ClaimAutoFillerBackend:
    def __init__(self, log_callback: Callable[[str], None]):
        self.log = log_callback
        self.amount_data = {} # 存储从 Excel 读取的 {保单号: 金额}

    def load_excel_data(self, excel_path: str):
        """读取 Excel 并建立索引"""
        if not excel_path or not os.path.exists(excel_path):
            return
        try:
            # 读取 Excel，强制将所有内容作为字符串处理，防止保单号变成科学计数法
            df = pd.read_excel(excel_path, dtype=str)
            # 清洗列名，去除空格
            df.columns = df.columns.str.strip()
            
            amount_col = "保险金额"
            if "保险金额" not in df.columns:
                if len(df.columns) >= 3:
                    amount_col = df.columns[2] # Fallback to Column C (0-based index 2)
                    self.log(f"⚠️ 未找到 '保险金额' 列，尝试使用第3列: '{amount_col}'")
                else:
                    self.log("❌ Excel 错误: 未找到 '保险金额' 且列数不足3列！")
                    return

            if "保险单号" not in df.columns:
                 self.log("❌ Excel 错误: 必须包含 '保险单号' 列！")
                 return

            # 建立字典: {去除空格的保单号 : 金额}
            self.amount_data = dict(zip(
                df["保险单号"].str.strip().str.replace(" ", ""), 
                df[amount_col].str.strip()
            ))
            self.log(f"📊 已加载 Excel 数据，共 {len(self.amount_data)} 条记录。")
        except Exception as e:
            self.log(f"❌ 读取 Excel 失败: {str(e)}")

    def parse_chinese_date(self, text: str) -> Optional[datetime]:
        match = re.search(r'(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*日', text)
        if match:
            year, month, day = match.groups()
            full_year = int(year) + 2000 if len(year) == 2 else int(year)
            return datetime(full_year, int(month), int(day))
        return None

    def get_info_from_source(self, doc: Document) -> Dict[str, str]:
        info = {}
        if len(doc.tables) > 0:
            source_table = doc.tables[0]
            for row in source_table.rows:
                cells = row.cells
                for i in range(len(cells) - 1):
                    text = cells[i].text.strip().replace(" ", "")
                    val_text = cells[i+1].text.strip()
                    if "保险单号" in text: info["policy_no"] = val_text
                    elif "被保险人" in text: info["insured"] = val_text
                    elif "标的" in text and "名称" not in text: info["subject"] = val_text
                    elif "报案号" in text: info["report_no"] = val_text
        
        for para in doc.paragraphs:
            if "抄单通知时间" in para.text:
                date_obj = self.parse_chinese_date(para.text)
                if date_obj: info["notice_date"] = date_obj
                break
        return info

    def generate_fill_data(self, info: Dict[str, str]) -> Dict[str, str]:
        fill_data = {}
        policy_no = info.get("policy_no", "")
        clean_policy_no = policy_no.replace(" ", "")
        
        fill_data["保险单号"] = policy_no
        fill_data["标的名称"] = info.get("subject", "")
        
        # 从 Excel 数据中查找金额
        if clean_policy_no in self.amount_data:
            fill_data["保险金额"] = self.amount_data[clean_policy_no]
        else:
            fill_data["保险金额"] = "【未在Excel中找到金额】"

        report_no = info.get("report_no", "")
        clean_report_no = report_no.replace("DSHH", "").strip()
        insurance_type = "未知险种"
        if len(clean_report_no) >= 6:
            type_code = clean_report_no[3:6]
            if type_code == "043": insurance_type = "水路货运险"
            elif type_code == "041": insurance_type = "公路货运险"
        fill_data["承保险种"] = insurance_type

        insured = info.get("insured", "")
        deductible = "请核实免赔条件"
        if "恒力石化" in insured:
            deductible = "2000元或损失金额的5%， 以高者为准；但最高不超过人民币200,000.00元"
        elif "浙江卓航多式联运科技有限公司" in insured:
            deductible = "参照协议CSHHHYX2024Q000427中约定"
        fill_data["免赔条件"] = deductible

        fill_data.update({
            "保费是否已付": "是", "是否共保（Y/N）": "N", "特约回分（Y/N）": "N",
            "是否临分（Y/N）": "N", "共保人 / 比例": "无", "特约回分联系人": "无",
            "临分联系人": "无"
        })

        if "notice_date" in info:
            next_day = info["notice_date"] + timedelta(days=1)
            fill_data["填 单 日 期"] = next_day.strftime("%Y-%m-%d")
        else:
            fill_data["填 单 日 期"] = datetime.now().strftime("%Y-%m-%d")
        return fill_data

    def _set_run_font(self, run):
        """设置字体为10号 (10pt), 中文楷体, 英文Arial"""
        run.font.size = Pt(10)
        run.font.name = 'Arial' # 英文/数字字体
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '楷体') # 中文字体

    def _insert_floating_picture(self, run, img_path, width=None, height=None):
        """插入浮动图片 (衬于文字下方, 居中)"""
        # 1. 插入内嵌图片
        pic = run.add_picture(img_path, width=width, height=height)
        
        # 2. 获取内嵌图片的 XML
        inline = pic._inline
        
        # 3. 构造 Anchor XML (浮动)
        # ID 和 Name 随便生成一个
        pic_id = "0" 
        pic_name = "Signature"
        
        # 核心 XML: wp:anchor with behindDoc="1"
        # 修改: positionH 使用 align=center 代替 posOffset
        # 修改: positionV 设置偏移量 -126000 (约 -0.35cm, 向上移动)
        anchor_xml = f"""
        <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="column">
            <wp:align>center</wp:align>
          </wp:positionH>
          <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>-126000</wp:posOffset>
          </wp:positionV>
          <wp:extent cx="{inline.extent.cx}" cy="{inline.extent.cy}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="{pic_id}" name="{pic_name}"/>
          <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
          </wp:cNvGraphicFramePr>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:nvPicPr>
                  <pic:cNvPr id="{pic_id}" name="{pic_name}"/>
                  <pic:cNvPicPr/>
                </pic:nvPicPr>
                <pic:blipFill>
                  <a:blip r:embed="{inline.graphic.graphicData.pic.blipFill.blip.embed}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
                  <a:stretch>
                    <a:fillRect/>
                  </a:stretch>
                </pic:blipFill>
                <pic:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="{inline.extent.cx}" cy="{inline.extent.cy}"/>
                  </a:xfrm>
                  <a:prstGeom prst="rect">
                    <a:avLst/>
                  </a:prstGeom>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
        """
        
        # 4. 解析并替换
        anchor_element = parse_xml(anchor_xml)
        inline.getparent().replace(inline, anchor_element)

    def _insert_signature_merged(self, cell, next_cell_below, mode: str, img_path: str):
        """插入大号签名 (不合并单元格，让图片浮动跨越表格线)"""
        # 注意: 这里不再执行 cell.merge(next_cell_below)
        
        # 清空当前单元格 (填单人)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中对齐
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 清空下方单元格 (复核人) - 保持空白以便图片显示
        # next_cell_below.text = "" # 可以选择不清空或者清空，根据需求，通常签名会盖住
        
        if mode == 'image':
            if not os.path.exists(img_path):
                self.log(f"⚠️ 找不到签名图片，回退文字。")
                run = paragraph.add_run("金磊")
                self._set_run_font(run)
                return
            try:
                run = paragraph.add_run()
                # 使用浮动图片插入 helper，指定固定尺寸 高2.2cm x 宽2.73cm
                self._insert_floating_picture(run, img_path, width=Cm(2.73), height=Cm(2.2))
            except Exception as e:
                self.log(f"❌ 图片失败: {e}")
                import traceback
                self.log(traceback.format_exc())
                run = paragraph.add_run("金磊")
                self._set_run_font(run)
        else:
            run = paragraph.add_run("金磊")
            self._set_run_font(run)

    def process(self, input_dir: str, output_dir: str, sign_mode: str, sign_path: str, excel_path: str):
        try:
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # 加载 Excel
            self.load_excel_data(excel_path)

            files = [f for f in os.listdir(input_dir) if f.endswith('.docx') and not f.startswith('~$')]
            self.log(f"📂 发现 {len(files)} 个 Word 文件...")

            success_count = 0
            for file_name in files:
                try:
                    full_path = os.path.join(input_dir, file_name)
                    doc = Document(full_path)
                    
                    info = self.get_info_from_source(doc)
                    if not info:
                        self.log(f"⚠️ 跳过 {file_name}")
                        continue
                    
                    fill_data = self.generate_fill_data(info)
                    
                    # 寻找目标表格
                    target_table_index = -1
                    for i, table in enumerate(doc.tables):
                        try:
                            if len(table.rows) > 0 and "保险单号" in table.rows[0].cells[0].text and i > 0:
                                target_table_index = i
                                break
                        except: continue
                    
                    if target_table_index != -1:
                        table = doc.tables[target_table_index]
                        rows = table.rows
                        # 使用 while 循环方便手动跳过行（当合并单元格后）
                        r_idx = 0
                        while r_idx < len(rows):
                            row = rows[r_idx]
                            skip_increment = False # 标记是否因合并而需要特殊处理索引
                            
                            for c in range(len(row.cells) - 1):
                                txt = row.cells[c].text.replace(" ", "").strip()
                                
                                # --- 处理签名 (特殊逻辑：合并单元格) ---
                                if "填单人" in txt:
                                    target_cell = row.cells[c+1]
                                    # 尝试获取下一行的对应单元格 (复核人右侧)
                                    if r_idx + 1 < len(rows):
                                        next_row_txt = rows[r_idx+1].cells[c].text.replace(" ", "").strip()
                                        # 确认下一行确实是 复核人
                                        if "复核人" in next_row_txt:
                                            next_cell = rows[r_idx+1].cells[c+1]
                                            self._insert_signature_merged(target_cell, next_cell, sign_mode, sign_path)
                                        else:
                                            # 如果下一行不是复核人，就只填当前格
                                            self._insert_signature_merged(target_cell, target_cell, sign_mode, sign_path)
                                    else:
                                        self._insert_signature_merged(target_cell, target_cell, sign_mode, sign_path)
                                    continue

                                # --- 处理常规文本 (五号字体) ---
                                for k, v in fill_data.items():
                                    if k.replace(" ", "") == txt:
                                        cell = row.cells[c+1]
                                        cell.text = "" # 清空
                                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER # 垂直居中
                                        paragraph = cell.paragraphs[0]
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT # 左对齐
                                        run = paragraph.add_run(str(v))
                                        self._set_run_font(run)

                            r_idx += 1
                        
                        save_path = os.path.join(output_dir, f"已填写_{file_name}")
                        doc.save(save_path)
                        self.log(f"✅ {file_name} | 金额: {fill_data.get('保险金额')}")
                        success_count += 1
                    else:
                        self.log(f"❌ 失败 {file_name}: 未找到表格")

                except Exception as e:
                    self.log(f"❌ 错误 {file_name}: {str(e)}")
            
            self.log(f"\n🎉 处理完成！共 {success_count} 个文件。")
            
        except Exception as e:
            self.log(f"❌ 初始化失败: {str(e)}")
            raise e # Re-raise so the outer thread wrapper catches it (though redundant, good for clarity)


# --- GUI 界面类 (前端) ---
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("保险自动抄单工具 V2.0")
        self.geometry("700x650") # 稍微调高一点以容纳新按钮
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(5, weight=1)

        # 1. 文件夹选择
        self.create_path_selector(0, "Word 输入文件夹:", "input_path_entry", "input_btn", is_folder=True)
        self.create_path_selector(1, "结果输出文件夹:", "output_path_entry", "output_btn", is_folder=True)
        
        # 2. Excel 数据源 (新增)
        self.create_path_selector(2, "Excel 数据表:", "excel_path_entry", "excel_btn", is_folder=False)

        # 3. 签名设置
        self.lbl_sign = ctk.CTkLabel(self, text="签名模式:", font=("Arial", 14, "bold"))
        self.lbl_sign.grid(row=3, column=0, padx=20, pady=(20, 10), sticky="w")

        self.sign_mode_var = tkinter.StringVar(value="text")
        self.radio_text = ctk.CTkRadioButton(self, text="文字签名", variable=self.sign_mode_var, value="text", command=self.toggle_sign_input)
        self.radio_text.grid(row=3, column=1, padx=10, pady=(20, 10), sticky="w")
        self.radio_img = ctk.CTkRadioButton(self, text="图片签名 (自动合并单元格)", variable=self.sign_mode_var, value="image", command=self.toggle_sign_input)
        self.radio_img.grid(row=3, column=1, padx=120, pady=(20, 10), sticky="w")

        self.sign_path_entry = ctk.CTkEntry(self, placeholder_text="签名图片路径...")
        self.sign_btn = ctk.CTkButton(self, text="选择图片", width=80, command=self.select_sign_img)

        # 4. 运行按钮
        self.btn_run = ctk.CTkButton(self, text="开始处理 (支持合并签名 + Excel匹配)", font=("Arial", 16, "bold"), height=40, command=self.start_processing_thread)
        self.btn_run.grid(row=4, column=0, columnspan=3, padx=20, pady=20, sticky="ew")

        # 5. 日志
        self.log_box = ctk.CTkTextbox(self, font=("Menlo", 12))
        self.log_box.grid(row=5, column=0, columnspan=3, padx=20, pady=(0, 20), sticky="nsew")
        self.log_box.insert("0.0", "请选择包含【保险单号】和【保险金额】列的 Excel 文件...\n")

        # 默认值
        self.input_path_entry.insert(0, os.path.expanduser("~/Documents"))
        self.output_path_entry.insert(0, os.path.expanduser("~/Documents/processed"))
    
    def create_path_selector(self, row, label_text, entry_attr, btn_attr, is_folder):
        label = ctk.CTkLabel(self, text=label_text, font=("Arial", 14))
        label.grid(row=row, column=0, padx=20, pady=10, sticky="w")
        entry = ctk.CTkEntry(self, placeholder_text="请选择路径...")
        entry.grid(row=row, column=1, padx=10, pady=10, sticky="ew")
        setattr(self, entry_attr, entry)
        cmd = lambda: self.browse(entry, is_folder)
        btn = ctk.CTkButton(self, text="浏览", width=80, command=cmd)
        btn.grid(row=row, column=2, padx=20, pady=10)
        setattr(self, btn_attr, btn)

    def browse(self, entry_widget, is_folder):
        if is_folder: path = filedialog.askdirectory()
        else: path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if path:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, path)

    def select_sign_img(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png *.jpg")])
        if file_path:
            self.sign_path_entry.delete(0, "end")
            self.sign_path_entry.insert(0, file_path)

    def toggle_sign_input(self):
        if self.sign_mode_var.get() == "image":
            self.sign_path_entry.grid(row=3, column=2, padx=20, pady=(20, 10), sticky="ew")
            if not self.sign_path_entry.get() and os.path.exists("签字.png"):
                 self.sign_path_entry.insert(0, os.path.abspath("签字.png"))
        else:
            self.sign_path_entry.grid_forget()

    def log(self, message):
        def _update():
            self.log_box.insert("end", message + "\n")
            self.log_box.see("end")
        self.after(0, _update)

    def start_processing_thread(self):
        input_dir = self.input_path_entry.get()
        output_dir = self.output_path_entry.get()
        excel_path = self.excel_path_entry.get()
        sign_mode = self.sign_mode_var.get()
        sign_path = self.sign_path_entry.get()

        if not input_dir or not output_dir:
            messagebox.showerror("错误", "请选择输入和输出文件夹！")
            return
        
        if not excel_path:
             messagebox.showwarning("提示", "未选择 Excel 文件，金额将无法自动填充！")

        if sign_mode == 'image' and not sign_path:
            messagebox.showwarning("提示", "请选择签名图片！")
            return

        self.btn_run.configure(state="disabled", text="处理中...")
        self.log_box.delete("0.0", "end")
        
        def run():
            try:
                backend = ClaimAutoFillerBackend(log_callback=self.log)
                backend.process(input_dir, output_dir, sign_mode, sign_path, excel_path)
            except Exception as e:
                self.log(f"💥 严重错误: 程序遇到未处理的异常: {e}")
                import traceback
                self.log(traceback.format_exc())
            finally:
                self.after(0, lambda: self.btn_run.configure(state="normal", text="开始处理 (支持合并签名 + Excel匹配)"))

        threading.Thread(target=run, daemon=True).start()

if __name__ == "__main__":
    app = App()
    app.mainloop()